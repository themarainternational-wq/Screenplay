"""
docx_export.py
--------------
Turns a list of (block_type, text) blocks into a properly formatted
.docx file - screenplay-style, not one giant paragraph.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(blocks, out_path, title=None):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(12)

    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for btype, text in blocks:
        if not text.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(text)

        if btype == "SCENE_HEADING":
            run.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif btype == "CHARACTER":
            p.paragraph_format.left_indent = Inches(2.0)
            p.paragraph_format.space_before = Pt(10)
            run.bold = True
        elif btype == "PARENTHETICAL":
            p.paragraph_format.left_indent = Inches(1.6)
            run.italic = True
        elif btype == "DIALOGUE":
            p.paragraph_format.left_indent = Inches(1.2)
            p.paragraph_format.right_indent = Inches(1.2)
        elif btype == "TRANSITION":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.bold = True
        else:  # ACTION
            p.paragraph_format.space_after = Pt(6)

    doc.save(out_path)
    return out_path
